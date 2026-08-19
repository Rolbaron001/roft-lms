#!/usr/bin/env python3
"""Build the ROFT Learning Dictionary as a Word document.

The document is generated from ``dictionary/lms-dictionary.json`` — the same
file the in-app dictionary page reads — so the circulated Word version and the
screen can never drift apart. Regenerate it whenever the JSON changes:

    python scripts/build-dictionary-docx.py

House style follows the ROFT documents already issued: Montserrat headings on
navy, Inter body text, gold accent.
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = "0D1E32"
GOLD = "B9975B"
INK = "1A1A1A"
PAPER = "F4F1EA"

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
DATA = REPO / "dictionary" / "lms-dictionary.json"
OUT_DIR = REPO.parent.parent / "Dictionary"

SOURCE_LABEL = {
    "authority": "Authority",
    "platform": "ROFT",
    "practice": "Practice",
}


def shade(element, fill: str) -> None:
    """Apply a solid background to a paragraph or a table cell."""
    properties = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_borders(cell, colour: str = "D8D3C8") -> None:
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "4")
        line.set(qn("w:color"), colour)
        borders.append(line)
    cell._tc.get_or_add_tcPr().append(borders)


def repeat_header(row) -> None:
    """Mark a table row as a header so Word repeats it across pages."""
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def text(
    paragraph,
    value: str,
    *,
    font: str = "Inter",
    size: int = 10.5,
    colour: str = INK,
    bold: bool = False,
    italic: bool = False,
):
    run = paragraph.add_run(value)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(colour)
    run.bold = bold
    run.italic = italic
    return run


def para(container, value="", *, space_before=0, space_after=6, align=None, **kwargs):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        paragraph.alignment = align
    if value:
        text(paragraph, value, **kwargs)
    return paragraph


def section_bar(document, number: int, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.left_indent = Cm(0.25)
    shade(paragraph._p, NAVY)
    text(
        paragraph,
        f"{number}   {title.upper()}",
        font="Montserrat",
        size=13,
        colour="FFFFFF",
        bold=True,
    )


def lead_in(document, label: str, body: str) -> None:
    paragraph = para(document, space_after=8)
    text(paragraph, f"{label}. ", font="Montserrat", size=10.5, colour=NAVY, bold=True)
    text(paragraph, body)


def build() -> Path:
    data = json.loads(DATA.read_text(encoding="utf-8"))
    entries = data["entries"]
    categories = data["categories"]

    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Inter"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    # ---------------------------------------------------------------- cover
    para(
        document,
        "ROFT STRATEGIC WORKFORCE ADVISORY",
        font="Montserrat",
        size=14,
        colour=NAVY,
        bold=True,
        space_after=2,
    )
    para(
        document,
        "Learning Programme Development",
        size=10.5,
        colour=GOLD,
        space_after=26,
    )

    banner = document.add_paragraph()
    banner.paragraph_format.space_before = Pt(14)
    banner.paragraph_format.space_after = Pt(4)
    banner.paragraph_format.left_indent = Cm(0.3)
    shade(banner._p, NAVY)
    text(
        banner,
        "THE ROFT LEARNING DICTIONARY",
        font="Montserrat",
        size=26,
        colour="FFFFFF",
        bold=True,
    )

    para(
        document,
        "Terms, abbreviations and concepts for occupational learning",
        font="Montserrat",
        size=13,
        colour=GOLD,
        space_after=16,
    )
    para(
        document,
        "The settled meaning of every term used across the ROFT Learning "
        "Management System, its documents and its learning material. Each "
        "entry says who owns the word, so that a term set by an authority is "
        "never confused with one we chose ourselves.",
        space_after=22,
    )

    para(
        document,
        "“Clarity. Structure. Better Decisions.”",
        font="Montserrat",
        size=13,
        colour=NAVY,
        bold=True,
        space_after=18,
    )

    for line, colour in [
        (f"Document reference: ROFT-LMS-DICT-{data['version']}", NAVY),
        (f"Version: {data['version']}", NAVY),
        (f"Issued: {format_date(data['issued'])}", NAVY),
        ("Status: In force", NAVY),
        ("Prepared by: ROFT Strategic Workforce Advisory", NAVY),
        ("roftbusiness.org", GOLD),
    ]:
        para(document, line, size=9.5, colour=colour, space_after=1)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------- 1 about this dictionary
    section_bar(document, 1, "About This Dictionary")

    lead_in(
        document,
        "Purpose",
        "This dictionary is the single reference for the terms, abbreviations "
        "and concepts used across ROFT's occupational learning work. It "
        "serves the Learning Management System, the programme design "
        "documents, and every piece of learning material written against a "
        "curriculum, so that a term means the same thing wherever it appears "
        "and anyone can find a plain explanation in one place.",
    )
    lead_in(
        document,
        "How it is arranged",
        "Entries are listed in a single alphabetical table. Abbreviations are "
        "given in brackets after the full term, so a reader can search on "
        "either. The right-hand column names who defines the term.",
    )
    lead_in(
        document,
        "Who defines it",
        "This is the column that matters, and it is why the dictionary exists. "
        "Authority means a body defines the term and we must use it as they "
        "do; changing the meaning puts a submission or an accreditation at "
        "risk, and the body is named. ROFT means we chose the word ourselves; "
        "it is ours to change and no regulator is watching. Practice means the "
        "term is widely used across the sector with no single owner; it is "
        "useful shorthand, but it must never be cited as a requirement.",
    )
    lead_in(
        document,
        "Standing",
        "The dictionary is a living reference, extended as new material "
        "introduces settled terms and corrected whenever a definition is "
        "refined at source. It does not create definitions of its own for "
        "terms an authority owns; it records the meaning already in use, in "
        "plain words. Where an entry is written in ROFT's own words rather "
        "than quoted, that is for consistency of voice, and the authority is "
        "named so the original can be checked.",
    )
    lead_in(
        document,
        "In the system",
        "The same entries are available to every signed-in user of the "
        "Learning Management System under Dictionary, searchable by term, by "
        "abbreviation and by any word in a definition. This document and that "
        "page are generated from one file, so they cannot disagree.",
    )

    start = document.add_paragraph()
    start.paragraph_format.space_before = Pt(14)
    start.paragraph_format.space_after = Pt(4)
    start.paragraph_format.left_indent = Cm(0.25)
    shade(start._p, PAPER)
    text(start, "START HERE", font="Montserrat", size=11, colour=GOLD, bold=True)

    body = document.add_paragraph()
    body.paragraph_format.space_after = Pt(10)
    body.paragraph_format.left_indent = Cm(0.25)
    shade(body._p, PAPER)
    text(
        body,
        "Four terms should be held before any other: Occupational "
        "Qualification, Curriculum Module, Internal Assessment Criterion and "
        "External Integrated Summative Assessment. A qualification is built "
        "around a job rather than a subject. It is delivered as modules of "
        "three kinds — knowledge, practical skill and work experience. "
        "Inside the knowledge and practical modules sit internal assessment "
        "criteria, every one of which the provider must judge achieved. Only "
        "then does the learner sit the external assessment, which is set and "
        "marked by somebody else. Everything else in this dictionary hangs "
        "off those four.",
    )

    counts = {"authority": 0, "platform": 0, "practice": 0}
    for entry in entries:
        counts[entry["definedBy"]] += 1
    para(
        document,
        f"{len(entries)} terms: {counts['authority']} set by an authority, "
        f"{counts['platform']} ROFT's own, {counts['practice']} common practice. "
        "Grouped across {} areas: {}.".format(
            len(categories),
            "; ".join(sorted(label.lower() for label in categories.values())),
        ),
        size=9.5,
        colour="5A6472",
        space_before=6,
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------------ 2 the dictionary
    section_bar(document, 2, "Dictionary of Terms")
    para(
        document,
        "Every term, abbreviation and concept in alphabetical order, with a "
        "short plain explanation of each and a note of who defines it.",
        space_after=10,
    )

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(4.6), Cm(9.6), Cm(2.0))

    header = table.rows[0]
    repeat_header(header)
    for cell, label in zip(header.cells, ("Term", "Explanation", "Defined by")):
        shade(cell._tc, NAVY)
        set_cell_borders(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.space_before = Pt(2)
        text(
            paragraph,
            label,
            font="Montserrat",
            size=9.5,
            colour="FFFFFF",
            bold=True,
        )

    for index, entry in enumerate(entries):
        row = table.add_row()
        stripe = PAPER if index % 2 else "FFFFFF"

        term_cell, explanation_cell, source_cell = row.cells
        for cell in row.cells:
            shade(cell._tc, stripe)
            set_cell_borders(cell)

        paragraph = term_cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        text(
            paragraph,
            entry["term"],
            font="Montserrat",
            size=9.5,
            colour=NAVY,
            bold=True,
        )
        if entry.get("abbreviation"):
            text(
                paragraph,
                f" ({entry['abbreviation']})",
                font="Montserrat",
                size=9.5,
                colour=GOLD,
                bold=True,
            )

        paragraph = explanation_cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        text(paragraph, entry["definition"], size=9.5)
        if entry.get("seeAlso"):
            text(
                paragraph,
                " See also " + ", ".join(entry["seeAlso"]) + ".",
                size=9.5,
                colour="5A6472",
                italic=True,
            )

        paragraph = source_cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        label = (
            entry["authority"]
            if entry["definedBy"] == "authority"
            else SOURCE_LABEL[entry["definedBy"]]
        )
        text(
            paragraph,
            label,
            font="Montserrat",
            size=8.5,
            colour=NAVY if entry["definedBy"] == "authority" else "5A6472",
            bold=entry["definedBy"] == "authority",
        )

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    para(
        document,
        "Sources consulted: the SAQA NQFpedia; the DHET Dictionary of Terms "
        "and Concepts for Post-School Education and Training; and the HSRC "
        "Labour Market Intelligence Partnership dictionary. Definitions here "
        "are written in ROFT's own words; where a term is owned by an "
        "authority, that authority is named so the original may be consulted.",
        size=9,
        colour="5A6472",
        space_before=12,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"ROFT Learning Dictionary V{data['version']}.docx"
    document.save(out)
    return out


def format_date(iso: str) -> str:
    parsed = date.fromisoformat(iso)
    return f"{parsed.day} {parsed.strftime('%B %Y')}"


if __name__ == "__main__":
    written = build()
    print(f"Wrote {written}")
